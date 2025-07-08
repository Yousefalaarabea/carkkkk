import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Data to include (copy from previous answer, as a Python dict/list structure)
SECTIONS = [
    {
        'title': '🔐 Authentication & Users',
        'endpoints': [
            {'method': 'POST', 'path': '/api/register/', 'desc': 'Register a new user', 'view': 'RegisterView', 'auth': 'No', 'params': 'JSON: email, password, phone_number, first_name, last_name, national_id', 'response': '201: user data or errors'},
            {'method': 'POST', 'path': '/api/login/', 'desc': 'Obtain JWT token', 'view': 'TokenObtainPairView', 'auth': 'No', 'params': 'JSON: email, password', 'response': '200: access/refresh tokens'},
            {'method': 'POST', 'path': '/api/token/refresh/', 'desc': 'Refresh JWT token', 'view': 'TokenRefreshView', 'auth': 'No', 'params': 'JSON: refresh token', 'response': '200: new access token'},
            {'method': 'GET', 'path': '/api/users/', 'desc': 'List all users (admin)', 'view': 'UserViewSet', 'auth': 'Admin Only', 'params': '-', 'response': '200: list of users'},
            {'method': 'POST', 'path': '/api/users/', 'desc': 'Create user (admin)', 'view': 'UserViewSet', 'auth': 'Admin Only', 'params': 'JSON: user fields', 'response': '201: user data'},
            {'method': 'GET', 'path': '/api/users/{id}/', 'desc': 'Retrieve user (admin)', 'view': 'UserViewSet', 'auth': 'Admin Only', 'params': '-', 'response': '200: user data'},
            {'method': 'PUT', 'path': '/api/users/{id}/', 'desc': 'Update user (admin)', 'view': 'UserViewSet', 'auth': 'Admin Only', 'params': 'JSON: user fields', 'response': '200: user data'},
            {'method': 'DELETE', 'path': '/api/users/{id}/', 'desc': 'Delete user (admin)', 'view': 'UserViewSet', 'auth': 'Admin Only', 'params': '-', 'response': '204: no content'},
            {'method': 'GET/POST/PUT/DELETE', 'path': '/api/roles/', 'desc': 'CRUD for roles', 'view': 'RoleViewSet', 'auth': 'No', 'params': 'JSON: role fields', 'response': '200/201/204'},
            {'method': 'GET/POST/PUT/DELETE', 'path': '/api/user-roles/', 'desc': 'CRUD for user roles', 'view': 'UserRoleViewSet', 'auth': 'Yes', 'params': 'JSON: user/role', 'response': '200/201/204'},
            {'method': 'POST', 'path': '/api/assign-roles/', 'desc': 'Assign roles to user', 'view': 'AssignRolesAPIView', 'auth': 'No', 'params': 'JSON: user_id, role_ids', 'response': '200: assigned roles'},
            {'method': 'GET', 'path': '/api/user-roles/{user_id}/', 'desc': 'Get roles for user', 'view': 'UserRolesAPIView', 'auth': 'No', 'params': '-', 'response': '200: roles list'},
        ]
    },
    {
        'title': '🚗 Cars Management',
        'endpoints': [
            {'method': 'GET', 'path': '/api/cars/', 'desc': 'List all cars', 'view': 'CarViewSet', 'auth': 'No', 'params': '-', 'response': '200: cars list'},
            {'method': 'POST', 'path': '/api/cars/', 'desc': 'Create a car', 'view': 'CarViewSet', 'auth': 'No', 'params': 'JSON: car fields', 'response': '201: car data'},
            {'method': 'GET', 'path': '/api/cars/{id}/', 'desc': 'Retrieve car', 'view': 'CarViewSet', 'auth': 'No', 'params': '-', 'response': '200: car data'},
            {'method': 'PUT', 'path': '/api/cars/{id}/', 'desc': 'Update car', 'view': 'CarViewSet', 'auth': 'No', 'params': 'JSON: car fields', 'response': '200: car data'},
            {'method': 'DELETE', 'path': '/api/cars/{id}/', 'desc': 'Delete car', 'view': 'CarViewSet', 'auth': 'No', 'params': '-', 'response': '204: no content'},
            {'method': 'GET', 'path': '/api/my-cars/', 'desc': 'List cars owned by user', 'view': 'MyCarsView', 'auth': 'Yes', 'params': '-', 'response': '200: cars list'},
            {'method': 'GET/POST/PUT/DELETE', 'path': '/api/car-rental-options/', 'desc': 'CRUD for rental options', 'view': 'CarRentalOptionsViewSet', 'auth': 'Yes', 'params': 'JSON: rental options', 'response': '200/201/204'},
            {'method': 'PATCH', 'path': '/api/car-rental-options/by-car/{car_id}/', 'desc': 'Update by car id', 'view': 'CarRentalOptionsViewSet', 'auth': 'Yes', 'params': 'JSON: fields', 'response': '200: updated data'},
            {'method': 'GET/POST/PUT/DELETE', 'path': '/api/car-usage-policy/', 'desc': 'CRUD for usage policy', 'view': 'CarUsagePolicyViewSet', 'auth': 'Yes', 'params': 'JSON: usage policy', 'response': '200/201/204'},
            {'method': 'PATCH', 'path': '/api/car-usage-policy/by-car/{car_id}/', 'desc': 'Update by car id', 'view': 'CarUsagePolicyViewSet', 'auth': 'Yes', 'params': 'JSON: fields', 'response': '200: updated data'},
            {'method': 'GET/POST/PUT/DELETE', 'path': '/api/car-stats/', 'desc': 'CRUD for car stats', 'view': 'CarStatsViewSet', 'auth': 'Yes', 'params': 'JSON: stats', 'response': '200/201/204'},
            {'method': 'PATCH', 'path': '/api/car-stats/by-car/{car_id}/', 'desc': 'Update by car id', 'view': 'CarStatsViewSet', 'auth': 'Yes', 'params': 'JSON: fields', 'response': '200: updated data'},
            {'method': 'GET', 'path': '/api/car-stats/by-car/{car_id}/', 'desc': 'Get stats by car id', 'view': 'CarStatsViewSet', 'auth': 'Yes', 'params': '-', 'response': '200: stats data'},
            {'method': 'GET', 'path': '/api/car-stats/summary/', 'desc': 'Get summary stats', 'view': 'CarStatsViewSet', 'auth': 'Yes', 'params': '-', 'response': '200: summary data'},
            {'method': 'POST', 'path': '/api/cars/test-validation/', 'desc': 'Validate car data (not saved)', 'view': 'CarValidationTestView', 'auth': 'Yes', 'params': 'JSON: stage, car data', 'response': '200: validation result'},
            {'method': 'POST', 'path': '/api/cars/test/basic-info/', 'desc': 'Test basic info validation', 'view': 'TestCarBasicInfoView', 'auth': 'Yes', 'params': 'JSON: fields', 'response': '200: result'},
            {'method': 'POST', 'path': '/api/cars/test/rental-options/', 'desc': 'Test rental options validation', 'view': 'TestCarRentalOptionsView', 'auth': 'Yes', 'params': 'JSON: fields', 'response': '200: result'},
            {'method': 'POST', 'path': '/api/cars/test/usage-policy/', 'desc': 'Test usage policy validation', 'view': 'TestCarUsagePolicyView', 'auth': 'Yes', 'params': 'JSON: fields', 'response': '200: result'},
            {'method': 'POST', 'path': '/api/cars/test/complete/', 'desc': 'Test complete car validation', 'view': 'TestCompleteCarView', 'auth': 'Yes', 'params': 'JSON: fields', 'response': '200: result'},
            {'method': 'POST', 'path': '/api/cars/test/plate-check/', 'desc': 'Quick plate number check', 'view': 'QuickPlateCheckView', 'auth': 'Yes', 'params': 'JSON: plate_number', 'response': '200: result'},
            {'method': 'POST', 'path': '/api/cars/test/pricing-suggestions/', 'desc': 'Get pricing suggestions', 'view': 'PricingSuggestionsView', 'auth': 'Yes', 'params': 'JSON: car_type, car_category, year', 'response': '200: suggestions'},
        ]
    },
    # ... (other sections as in previous answer)
]

def generate_word(filename):
    doc = Document()
    doc.add_heading('API Endpoints Summary', 0)
    for section in SECTIONS:
        doc.add_heading(section['title'], level=1)
        table = doc.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Method'
        hdr_cells[1].text = 'Path'
        hdr_cells[2].text = 'Description'
        hdr_cells[3].text = 'View/Class'
        hdr_cells[4].text = 'Auth Required'
        hdr_cells[5].text = 'Request Params'
        hdr_cells[6].text = 'Response'
        for ep in section['endpoints']:
            row_cells = table.add_row().cells
            row_cells[0].text = ep['method']
            row_cells[1].text = ep['path']
            row_cells[2].text = ep['desc']
            row_cells[3].text = ep['view']
            row_cells[4].text = ep['auth']
            row_cells[5].text = ep['params']
            row_cells[6].text = ep['response']
        doc.add_paragraph('')
    doc.save(filename)

def generate_pdf(filename):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    y = height - 40
    c.setFont('Helvetica-Bold', 16)
    c.drawString(40, y, 'API Endpoints Summary')
    y -= 30
    c.setFont('Helvetica', 12)
    for section in SECTIONS:
        c.setFont('Helvetica-Bold', 14)
        c.drawString(40, y, section['title'])
        y -= 20
        c.setFont('Helvetica', 10)
        headers = ['Method', 'Path', 'Description', 'View/Class', 'Auth', 'Params', 'Response']
        c.drawString(40, y, ' | '.join(headers))
        y -= 15
        for ep in section['endpoints']:
            line = f"{ep['method']} | {ep['path']} | {ep['desc']} | {ep['view']} | {ep['auth']} | {ep['params']} | {ep['response']}"
            c.drawString(40, y, line[:120])  # truncate for width
            y -= 12
            if y < 60:
                c.showPage()
                y = height - 40
        y -= 10
    c.save()

if __name__ == '__main__':
    generate_word('API_Endpoints_Summary.docx')
    generate_pdf('API_Endpoints_Summary.pdf')
    print('Documents generated: API_Endpoints_Summary.docx, API_Endpoints_Summary.pdf') 